from __future__ import annotations

import html
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from lxml import html as lxml_html
from pypdf import PdfReader

try:
    import openpyxl
except Exception:  # pragma: no cover
    openpyxl = None


ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT = ROOT.parent
OUT_DIR = ROOT / "documentacao" / "acervo"
SUPPORTED = {".docx", ".pdf", ".html", ".xlsx", ".png"}


def clean_text(value: str) -> str:
    value = html.unescape(value)
    value = value.replace("\x00", "")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def read_docx(path: Path) -> str:
    doc = Document(path)
    parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return clean_text("\n".join(parts))


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages[:12]:
        parts.append(page.extract_text() or "")
    return clean_text("\n".join(parts))


def read_html(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    doc = lxml_html.fromstring(raw)
    return clean_text(doc.text_content())


def read_xlsx(path: Path) -> str:
    if openpyxl is None:
        return ""
    workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
    rows: list[str] = []
    for sheet in workbook.worksheets[:4]:
        rows.append(f"Planilha: {sheet.title}")
        for row in sheet.iter_rows(max_row=40, values_only=True):
            cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if cells:
                rows.append(" | ".join(cells))
    return clean_text("\n".join(rows))


def read_file(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            return read_docx(path)
        if suffix == ".pdf":
            return read_pdf(path)
        if suffix == ".html":
            return read_html(path)
        if suffix == ".xlsx":
            return read_xlsx(path)
    except Exception as exc:
        return f"[Erro ao extrair texto: {exc}]"
    return ""


def category_for(path: Path) -> str:
    raw = str(path).lower()
    if "logomarca" in raw or path.suffix.lower() == ".png":
        return "Marca e assets"
    if "whatsapp" in raw or "vendas" in raw or "oferta" in raw:
        return "Vendas e lançamento"
    if "ipc" in raw or "icp" in raw:
        return "ICP e personas"
    if "live" in raw or "reels" in raw or "hooks" in raw or "carrossel" in raw:
        return "Conteúdo"
    if "metodo" in raw or "metodologia" in raw or "manual" in raw or "rotinas" in raw:
        return "Método e operação"
    return "Estratégia"


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    files = [
        path
        for path in PROJECT_ROOT.rglob("*")
        if path.is_file()
        and path.suffix.lower() in SUPPORTED
        and "/.git/" not in str(path)
        and not path.name.startswith("._")
        and "node_modules" not in path.parts
        and ".next" not in path.parts
    ]
    files.sort(key=lambda item: str(item).lower())

    inventory_lines = [
        "# Inventário do Acervo Triade Saúde e Performance",
        "",
        f"Atualizado em: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"Pasta analisada: `{PROJECT_ROOT}`",
        f"Total de arquivos relevantes: {len(files)}",
        "",
    ]

    extracted_lines = [
        "# Extração Textual do Acervo",
        "",
        "Este arquivo consolida trechos extraídos dos documentos locais para orientar o painel, o brandbook e os documentos doutrinários.",
        "",
    ]

    for path in files:
        rel = path.relative_to(PROJECT_ROOT)
        category = category_for(path)
        size_kb = path.stat().st_size / 1024
        inventory_lines.append(f"- **{category}** | `{rel}` | {path.suffix.lower()} | {size_kb:.1f} KB")
        text = read_file(path)
        if text:
            preview = text[:5000]
            extracted_lines.extend(
                [
                    f"## {rel}",
                    "",
                    f"Categoria: **{category}**",
                    "",
                    "```text",
                    preview,
                    "```",
                    "",
                ]
            )

    (OUT_DIR / "INVENTARIO.md").write_text("\n".join(inventory_lines) + "\n", encoding="utf-8")
    (OUT_DIR / "EXTRACAO-TEXTUAL.md").write_text("\n".join(extracted_lines) + "\n", encoding="utf-8")
    print(f"Inventário salvo em {OUT_DIR / 'INVENTARIO.md'}")
    print(f"Extração salva em {OUT_DIR / 'EXTRACAO-TEXTUAL.md'}")


if __name__ == "__main__":
    main()
